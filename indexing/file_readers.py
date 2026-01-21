"""
indexing/file_readers.py
File readers for extracting text from different file types
"""

from pathlib import Path
from typing import Optional
from abc import ABC, abstractmethod

from utils.logger import get_logger
from utils.file_utils import read_text_file, get_file_size

logger = get_logger(__name__)


class FileReadError(Exception):
    """Exception raised when file reading fails"""
    pass


class BaseFileReader(ABC):
    """Abstract base class for file readers"""
    
    @abstractmethod
    def read(self, filepath: Path) -> str:
        """
        Read file and extract text
        
        Args:
            filepath: Path to file
        
        Returns:
            Extracted text as string
        
        Raises:
            FileReadError: If reading fails
        """
        pass
    
    def can_read(self, filepath: Path) -> bool:
        """
        Check if this reader can read the file
        
        Args:
            filepath: Path to file
        
        Returns:
            True if can read, False otherwise
        """
        return filepath.suffix.lower() in self.supported_extensions
    
    @property
    @abstractmethod
    def supported_extensions(self) -> list:
        """List of supported file extensions"""
        pass


class TextFileReader(BaseFileReader):
    """Reader for plain text files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.txt', '.md', '.asc']
    
    def read(self, filepath: Path) -> str:
        """Read plain text file"""
        try:
            # Try UTF-8 first
            content = read_text_file(filepath, encoding='utf-8', errors='strict')
            if content is not None:
                logger.debug(f"Read text file (UTF-8): {filepath}")
                return content
            
            # Fallback to UTF-8 with error replacement
            content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if content is not None:
                logger.warning(f"Read text file with encoding errors replaced: {filepath}")
                return content
            
            # Last resort: latin-1 (never fails)
            content = read_text_file(filepath, encoding='latin-1', errors='replace')
            if content is not None:
                logger.warning(f"Read text file with latin-1 encoding: {filepath}")
                return content
            
            raise FileReadError(f"Failed to read text file: {filepath}")
        
        except Exception as e:
            logger.error(f"Error reading text file {filepath}: {e}")
            raise FileReadError(f"Failed to read text file: {str(e)}")


class MarkdownFileReader(BaseFileReader):
    """Reader for Markdown files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.md', '.markdown']
    
    def read(self, filepath: Path) -> str:
        """Read Markdown file (same as text, but could add parsing later)"""
        try:
            import markdown
            
            # Read raw markdown
            content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if not content:
                raise FileReadError(f"Failed to read markdown file: {filepath}")
            
            # Convert to plain text (strip markdown formatting)
            # For now, just return raw markdown as it's still readable
            # Could use markdown.markdown() to convert to HTML then strip tags
            logger.debug(f"Read markdown file: {filepath}")
            return content
        
        except ImportError:
            # Fallback if markdown library not installed
            logger.warning("markdown library not installed, reading as plain text")
            return TextFileReader().read(filepath)
        
        except Exception as e:
            logger.error(f"Error reading markdown file {filepath}: {e}")
            raise FileReadError(f"Failed to read markdown file: {str(e)}")


class PDFFileReader(BaseFileReader):
    """Reader for PDF files with OCR support for image-based PDFs"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.pdf']
    
    def read(self, filepath: Path) -> str:
        """Read PDF file and extract text (with OCR fallback)"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(str(filepath))
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num} in {filepath}: {e}")
                    continue
            
            # If we got text, return it
            if text_parts:
                content = "\n\n".join(text_parts)
                logger.debug(f"Read PDF file ({len(reader.pages)} pages): {filepath}")
                return content
            
            # No text extracted - try OCR
            logger.info(f"No text in PDF, attempting OCR: {filepath}")
            return self._ocr_pdf(filepath)
        
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise FileReadError("PyPDF2 library not available")
        
        except Exception as e:
            logger.error(f"Error reading PDF file {filepath}: {e}")
            raise FileReadError(f"Failed to read PDF file: {str(e)}")
    
    def _ocr_pdf(self, filepath: Path) -> str:
        """
        Extract text from image-based PDF using OCR
        
        Args:
            filepath: Path to PDF file
        
        Returns:
            Extracted text
        
        Raises:
            FileReadError: If OCR fails or libraries not available
        """
        try:
            import pytesseract
            from pdf2image import convert_from_path
            
            logger.info(f"Starting OCR on PDF: {filepath}")
            
            # Convert PDF pages to images
            try:
                images = convert_from_path(str(filepath), dpi=300)  # Higher DPI for better OCR
            except Exception as e:
                logger.error(f"Failed to convert PDF to images: {e}")
                raise FileReadError(f"Failed to convert PDF to images: {str(e)}")
            
            if not images:
                raise FileReadError("PDF conversion resulted in no images")
            
            # OCR each page
            text_parts = []
            total_pages = len(images)
            
            for i, image in enumerate(images):
                try:
                    logger.debug(f"OCR processing page {i+1}/{total_pages}")
                    
                    # Perform OCR
                    text = pytesseract.image_to_string(image, lang='eng')
                    
                    if text.strip():
                        text_parts.append(text.strip())
                    else:
                        logger.warning(f"No text detected on page {i+1}")
                
                except Exception as e:
                    logger.warning(f"OCR failed for page {i+1}: {e}")
                    continue
            
            if not text_parts:
                raise FileReadError(
                    f"No text extracted from PDF (even with OCR): {filepath}. "
                    "The PDF may be empty or the images may be unreadable."
                )
            
            content = "\n\n".join(text_parts)
            logger.info(f"OCR completed: {filepath} ({total_pages} pages, {len(text_parts)} with text)")
            return content
        
        except ImportError as e:
            missing_lib = "pytesseract or pdf2image"
            if "pytesseract" in str(e):
                missing_lib = "pytesseract"
            elif "pdf2image" in str(e):
                missing_lib = "pdf2image"
            
            error_msg = (
                f"OCR libraries not installed ({missing_lib} missing). "
                f"Install with:\n"
                f"  pip install pytesseract pdf2image\n"
                f"  brew install tesseract poppler  # macOS\n"
                f"  # or on Linux: sudo apt-get install tesseract-ocr poppler-utils"
            )
            logger.warning(error_msg)
            
            raise FileReadError(
                f"PDF contains only images and OCR is not available. "
                f"Install OCR support: pip install pytesseract pdf2image && "
                f"brew install tesseract poppler"
            )
        
        except FileReadError:
            # Re-raise FileReadError as-is
            raise
        
        except Exception as e:
            logger.error(f"OCR failed for {filepath}: {e}")
            raise FileReadError(f"Failed to OCR PDF: {str(e)}")


class DocxFileReader(BaseFileReader):
    """Reader for Microsoft Word documents (.docx)"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.docx']
    
    def read(self, filepath: Path) -> str:
        """Read .docx file and extract text"""
        try:
            from docx import Document
            
            doc = Document(str(filepath))
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            if not text_parts:
                raise FileReadError(f"No text extracted from DOCX: {filepath}")
            
            content = "\n\n".join(text_parts)
            logger.debug(f"Read DOCX file: {filepath}")
            return content
        
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise FileReadError("python-docx library not available")
        
        except Exception as e:
            logger.error(f"Error reading DOCX file {filepath}: {e}")
            raise FileReadError(f"Failed to read DOCX file: {str(e)}")


class HTMLFileReader(BaseFileReader):
    """Reader for HTML files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.html', '.htm']
    
    def read(self, filepath: Path) -> str:
        """Read HTML file and extract text"""
        try:
            from bs4 import BeautifulSoup
            
            # Read HTML content
            html_content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if not html_content:
                raise FileReadError(f"Failed to read HTML file: {filepath}")
            
            # Parse and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                raise FileReadError(f"No text extracted from HTML: {filepath}")
            
            logger.debug(f"Read HTML file: {filepath}")
            return text
        
        except ImportError:
            logger.error("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            raise FileReadError("beautifulsoup4 library not available")
        
        except Exception as e:
            logger.error(f"Error reading HTML file {filepath}: {e}")
            raise FileReadError(f"Failed to read HTML file: {str(e)}")


class CSVFileReader(BaseFileReader):
    """Reader for CSV files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.csv', '.tsv']
    
    def read(self, filepath: Path) -> str:
        """Read CSV file and convert to text"""
        try:
            import csv
            
            # Detect delimiter
            delimiter = ',' if filepath.suffix.lower() == '.csv' else '\t'
            
            text_parts = []
            
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f, delimiter=delimiter)
                
                for row_num, row in enumerate(reader):
                    # Join cells with | separator
                    row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                    if row_text:
                        text_parts.append(row_text)
                    
                    # Limit to first 1000 rows for very large CSVs
                    if row_num >= 1000:
                        logger.warning(f"CSV file too large, truncated to 1000 rows: {filepath}")
                        break
            
            if not text_parts:
                raise FileReadError(f"No data extracted from CSV: {filepath}")
            
            content = "\n".join(text_parts)
            logger.debug(f"Read CSV file ({len(text_parts)} rows): {filepath}")
            return content
        
        except Exception as e:
            logger.error(f"Error reading CSV file {filepath}: {e}")
            raise FileReadError(f"Failed to read CSV file: {str(e)}")


class CodeFileReader(BaseFileReader):
    """Reader for code files (Python, Java, JavaScript, etc.)"""
    
    @property
    def supported_extensions(self) -> list:
        return [
            '.py', '.java', '.js', '.jsx', '.ts', '.tsx',
            '.cpp', '.c', '.h', '.hpp', '.cs', '.rb',
            '.go', '.rs', '.php', '.swift', '.kt',
            '.css', '.json', '.yaml', '.yml', '.xml'
        ]
    
    def read(self, filepath: Path) -> str:
        """Read code file (plain text with UTF-8 encoding)"""
        try:
            # Try UTF-8 first (most common for code)
            content = read_text_file(filepath, encoding='utf-8', errors='strict')
            if content is not None:
                logger.debug(f"Read code file: {filepath}")
                return content
            
            # Fallback with error replacement
            content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if content is not None:
                logger.warning(f"Read code file with encoding errors replaced: {filepath}")
                return content
            
            raise FileReadError(f"Failed to read code file: {filepath}")
        
        except Exception as e:
            logger.error(f"Error reading code file {filepath}: {e}")
            raise FileReadError(f"Failed to read code file: {str(e)}")


class ConfigFileReader(BaseFileReader):
    """Reader for configuration files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.ini', '.cfg', '.conf', '.toml', '.log']
    
    def read(self, filepath: Path) -> str:
        """Read configuration file (plain text)"""
        try:
            content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if not content:
                raise FileReadError(f"Failed to read config file: {filepath}")
            
            logger.debug(f"Read config file: {filepath}")
            return content
        
        except Exception as e:
            logger.error(f"Error reading config file {filepath}: {e}")
            raise FileReadError(f"Failed to read config file: {str(e)}")


class RTFFileReader(BaseFileReader):
    """Reader for RTF (Rich Text Format) files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.rtf']
    
    def read(self, filepath: Path) -> str:
        """Read RTF file - basic implementation"""
        try:
            # RTF is complex format, for now read as text and strip RTF codes
            content = read_text_file(filepath, encoding='utf-8', errors='replace')
            if not content:
                raise FileReadError(f"Failed to read RTF file: {filepath}")
            
            # Basic RTF code stripping (not perfect but functional)
            import re
            
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*\s?', ' ', content)
            # Remove braces
            content = content.replace('{', '').replace('}', '')
            # Remove backslash escapes
            content = content.replace('\\', '')
            # Clean up whitespace
            content = re.sub(r'\s+', ' ', content)
            
            if not content.strip():
                raise FileReadError(f"No text extracted from RTF: {filepath}")
            
            logger.debug(f"Read RTF file (basic extraction): {filepath}")
            logger.warning("RTF support is basic. For better results, convert to DOCX or PDF")
            return content.strip()
        
        except Exception as e:
            logger.error(f"Error reading RTF file {filepath}: {e}")
            raise FileReadError(f"Failed to read RTF file: {str(e)}")


class DocFileReader(BaseFileReader):
    """Reader for Microsoft Word 97-2003 documents (.doc)"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.doc']
    
    def read(self, filepath: Path) -> str:
        """
        Read .doc file and extract text
        
        Note: .doc files are binary format and require antiword or similar tools.
        This implementation tries multiple approaches.
        """
        try:
            # Try using python-docx (sometimes works for .doc files)
            try:
                from docx import Document
                doc = Document(str(filepath))
                
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                if text_parts:
                    content = "\n\n".join(text_parts)
                    logger.debug(f"Read DOC file with python-docx: {filepath}")
                    return content
            except Exception:
                pass  # Try next method
            
            # Try using textract library (if available)
            try:
                import textract
                text = textract.process(str(filepath)).decode('utf-8')
                if text.strip():
                    logger.debug(f"Read DOC file with textract: {filepath}")
                    return text
            except ImportError:
                logger.warning("textract not installed. Install with: pip install textract")
            except Exception as e:
                logger.debug(f"textract failed for {filepath}: {e}")
            
            # Try using antiword (command-line tool)
            try:
                import subprocess
                result = subprocess.run(
                    ['antiword', str(filepath)],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0 and result.stdout.strip():
                    logger.debug(f"Read DOC file with antiword: {filepath}")
                    return result.stdout
            except FileNotFoundError:
                logger.warning("antiword not found. Install with: brew install antiword (macOS)")
            except Exception as e:
                logger.debug(f"antiword failed for {filepath}: {e}")
            
            # Last resort: try reading as text (often produces garbage but better than nothing)
            logger.warning(f"Attempting to read .doc as plain text (may produce poor results): {filepath}")
            content = read_text_file(filepath, encoding='latin-1', errors='ignore')
            if content and len(content.strip()) > 100:  # At least some text extracted
                return content
            
            raise FileReadError(
                "Failed to read .doc file. Please convert to .docx format for better results. "
                "You can install antiword (brew install antiword) or textract (pip install textract) "
                "for better .doc support."
            )
        
        except Exception as e:
            logger.error(f"Error reading DOC file {filepath}: {e}")
            raise FileReadError(f"Failed to read .doc file: {str(e)}")


class PagesFileReader(BaseFileReader):
    """Reader for Apple Pages documents (.pages)"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.pages']
    
    def read(self, filepath: Path) -> str:
        """
        Read .pages file and extract text
        
        Note: .pages files are actually ZIP archives containing XML/PDF
        """
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            
            # .pages files are ZIP archives
            if not zipfile.is_zipfile(filepath):
                raise FileReadError(f"Invalid .pages file (not a ZIP archive): {filepath}")
            
            text_parts = []
            
            with zipfile.ZipFile(filepath, 'r') as zip_ref:
                # Try to find and read index.xml (contains text content)
                try:
                    # Look for the main content file
                    content_files = [
                        'index.xml',
                        'Index/Document.iwa',
                        'QuickLook/Preview.pdf'
                    ]
                    
                    # Try index.xml first (older .pages format)
                    if 'index.xml' in zip_ref.namelist():
                        with zip_ref.open('index.xml') as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            
                            # Extract text from XML (simplified extraction)
                            for elem in root.iter():
                                if elem.text and elem.text.strip():
                                    text_parts.append(elem.text.strip())
                        
                        if text_parts:
                            content = "\n".join(text_parts)
                            logger.debug(f"Read Pages file from index.xml: {filepath}")
                            return content
                
                except Exception as e:
                    logger.debug(f"Failed to read index.xml from {filepath}: {e}")
                
                # Try QuickLook preview PDF (newer .pages format)
                try:
                    if 'QuickLook/Preview.pdf' in zip_ref.namelist():
                        # Extract PDF and read it
                        with zip_ref.open('QuickLook/Preview.pdf') as pdf_file:
                            import io
                            from PyPDF2 import PdfReader
                            
                            pdf_bytes = io.BytesIO(pdf_file.read())
                            pdf_reader = PdfReader(pdf_bytes)
                            
                            for page in pdf_reader.pages:
                                text = page.extract_text()
                                if text:
                                    text_parts.append(text)
                            
                            if text_parts:
                                content = "\n\n".join(text_parts)
                                logger.debug(f"Read Pages file from Preview.pdf: {filepath}")
                                return content
                
                except ImportError:
                    logger.warning("PyPDF2 not installed, cannot read Preview.pdf from .pages file")
                except Exception as e:
                    logger.debug(f"Failed to read Preview.pdf from {filepath}: {e}")
                
                # Try textract as fallback (if available)
                try:
                    import textract
                    text = textract.process(str(filepath)).decode('utf-8')
                    if text.strip():
                        logger.debug(f"Read Pages file with textract: {filepath}")
                        return text
                except ImportError:
                    pass
                except Exception as e:
                    logger.debug(f"textract failed for {filepath}: {e}")
            
            raise FileReadError(
                "Failed to extract text from .pages file. "
                "Please export to .docx or .pdf format for better results."
            )
        
        except zipfile.BadZipFile:
            logger.error(f"Invalid .pages file (corrupted ZIP): {filepath}")
            raise FileReadError("Invalid or corrupted .pages file")
        
        except Exception as e:
            logger.error(f"Error reading Pages file {filepath}: {e}")
            raise FileReadError(f"Failed to read .pages file: {str(e)}")
# ============================================================================
# Factory Function
# ============================================================================

# Registry of readers
_READERS = [
    TextFileReader(),
    MarkdownFileReader(),
    PDFFileReader(),
    DocxFileReader(),
    DocFileReader(),
    PagesFileReader(),
    HTMLFileReader(),
    CSVFileReader(),
    CodeFileReader(),
    ConfigFileReader(),
    RTFFileReader(),
]


def get_reader_for_file(filepath: Path) -> Optional[BaseFileReader]:
    """
    Get appropriate reader for file based on extension
    
    Args:
        filepath: Path to file
    
    Returns:
        FileReader instance or None if no reader found
    """
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    
    for reader in _READERS:
        if extension in reader.supported_extensions:
            return reader
    
    logger.warning(f"No reader found for file type: {extension}")
    return None


def read_file(filepath: Path) -> str:
    """
    Read file and extract text (convenience function)
    
    Args:
        filepath: Path to file
    
    Returns:
        Extracted text
    
    Raises:
        FileReadError: If no reader found or reading fails
    """
    filepath = Path(filepath)
    
    reader = get_reader_for_file(filepath)
    if not reader:
        raise FileReadError(f"No reader available for file type: {filepath.suffix}")
    
    return reader.read(filepath)


def get_supported_extensions() -> list:
    """
    Get list of all supported file extensions
    
    Returns:
        List of extensions (e.g., ['.txt', '.pdf', ...])
    """
    extensions = set()
    for reader in _READERS:
        extensions.update(reader.supported_extensions)
    
    return sorted(list(extensions))


# ============================================================================
# Testing/Debug
# ============================================================================

def test_readers():
    """Test all readers (for development)"""
    logger.info("Testing file readers...")
    
    for reader in _READERS:
        logger.info(f"{reader.__class__.__name__}: {reader.supported_extensions}")
    
    logger.info(f"Total supported extensions: {len(get_supported_extensions())}")
    logger.info(f"Supported extensions: {get_supported_extensions()}")


if __name__ == "__main__":
    test_readers()